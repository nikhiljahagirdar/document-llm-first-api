from fastapi import APIRouter, Depends, HTTPException, status
import psycopg
from app.db_raw import get_raw_db
from app.schemas import TemplateResponse, TemplateCreate, TemplateUpdate, TemplateGenerateRequest
from typing import List, Optional, Any
import uuid
import json
from datetime import datetime
from app.dependencies import get_current_user
from app.services.render_service import generate_docx_template, render_html_template, parse_csv_data
from app.services.storage_service import upload_bytes_to_s3
from app.services.llm_service import LLMService
from app.services.db.template_db_service import TemplateDBService
from app.services.db.industry_db_service import IndustryDBService
from app.services.db.category_db_service import CategoryDBService
from app.services.db.subcategory_db_service import SubcategoryDBService
from app.services.db.audit_log_db_service import AuditLogDBService

router = APIRouter(prefix="/templates", tags=["templates"])

async def get_template_service():
    return TemplateDBService()

async def get_industry_service():
    return IndustryDBService()

async def get_category_service():
    return CategoryDBService()

async def get_subcategory_service():
    return SubcategoryDBService()

@router.get("/public/", response_model=List[TemplateResponse])
async def get_public_templates(
    search: Optional[str] = None,
    limit: int = 20,
    offset: int = 0,
    industry_id: Optional[uuid.UUID] = None,
    category_id: Optional[uuid.UUID] = None,
    subcategory_id: Optional[uuid.UUID] = None,
    conn: psycopg.AsyncConnection = Depends(get_raw_db),
    service: TemplateDBService = Depends(get_template_service)
):
    """
    Retrieve all public document generation templates with search and pagination.
    """
    return await service.list_public_templates(conn, limit, offset, industry_id, category_id, subcategory_id, search)


@router.get("/my/", response_model=List[TemplateResponse])
async def get_my_templates(
    conn: psycopg.AsyncConnection = Depends(get_raw_db),
    current_user: Any = Depends(get_current_user),
    service: TemplateDBService = Depends(get_template_service)
):
    """
    Retrieve templates owned by the current tenant or public templates.
    """
    return await service.list_user_templates(conn, current_user.tenant_id)


@router.post("/ai-builder", response_model=TemplateResponse)
async def ai_template_builder(
    industry_id: uuid.UUID,
    category_id: uuid.UUID,
    subcategory_id: Optional[uuid.UUID] = None,
    conn: psycopg.AsyncConnection = Depends(get_raw_db),
    current_user: Any = Depends(get_current_user),
    service: TemplateDBService = Depends(get_template_service),
    ind_service: IndustryDBService = Depends(get_industry_service),
    cat_service: CategoryDBService = Depends(get_category_service),
    sub_service: SubcategoryDBService = Depends(get_subcategory_service)
):
    """
    AI Template Builder: Generates a complete professional template based on classification.
    """
    # Fetch names
    ind = await ind_service.get_industry(conn, industry_id)
    cat = await cat_service.get_category(conn, category_id)
    sub = None
    if subcategory_id:
        sub = await sub_service.get_subcategory(conn, subcategory_id)

    industry_name = ind["name"] if ind else "N/A"
    category_name = cat["name"] if cat else "N/A"
    subcategory_name = sub["name"] if sub else "N/A"
    subcategory_prompt = sub.get("prompt") if sub else None

    ai_template = await LLMService.generate_template_ai(
        industry_name, category_name, subcategory_name, current_user.tenant_id, conn, override_prompt=subcategory_prompt
    )

    if "error" in ai_template:
        raise HTTPException(status_code=500, detail=ai_template["error"])

    now = datetime.now()
    template_data = {
        "template_name": ai_template.get("template_name"),
        "description": ai_template.get("description"),
        "template_schema": ai_template.get("template_schema"),
        "html_content": ai_template.get("html_content"),
        "document_type": ai_template.get("document_type"),
        "industry_id": industry_id,
        "category_id": category_id,
        "subcategory_id": subcategory_id,
        "category": category_name,
        "subcategory": subcategory_name,
        "tenant_id": current_user.tenant_id,
        "user_id": current_user.user_id,
        "is_public": False,
        "created_on": now,
        "updated_on": now
    }
    
    result = await service.create_template(conn, template_data)
    
    # Record Audit Log
    await AuditLogDBService.record_audit_log(
        conn, current_user.tenant_id, current_user.user_id,
        "template_ai_build", "template", str(result["template_id"]),
        {"template_name": result["template_name"]}
    )
    
    return result


@router.post("/", response_model=TemplateResponse)
async def create_template(
    template_data: TemplateCreate, 
    conn: psycopg.AsyncConnection = Depends(get_raw_db),
    current_user: Any = Depends(get_current_user),
    service: TemplateDBService = Depends(get_template_service),
    ind_service: IndustryDBService = Depends(get_industry_service),
    cat_service: CategoryDBService = Depends(get_category_service),
    sub_service: SubcategoryDBService = Depends(get_subcategory_service)
):
    industry_name = "N/A"
    category_name = "N/A"
    subcategory_name = "N/A"

    # Task: If industry/category/subcategory are not provided, use AI to generate the template
    if not (template_data.industry_id or template_data.category_id or template_data.subcategory_id) and template_data.prompt:
        ai_template = await LLMService.generate_template_ai(
            None, None, None, current_user.tenant_id, conn, user_prompt=template_data.prompt
        )
        
        if "error" in ai_template:
            raise HTTPException(status_code=500, detail=ai_template["error"])
            
        # Update template_data with AI results
        template_data.template_name = template_data.template_name or ai_template.get("template_name", "AI Generated Template")
        template_data.description = ai_template.get("description")
        template_data.template_schema = ai_template.get("template_schema")
        template_data.html_content = ai_template.get("html_content")
        template_data.document_type = ai_template.get("document_type")

    # Proceed with normal creation/rendering
    if template_data.industry_id:
        ind = await ind_service.get_industry(conn, template_data.industry_id)
        if ind: industry_name = ind["name"]

    if template_data.category_id:
        cat = await cat_service.get_category(conn, template_data.category_id)
        if cat: category_name = cat["name"]

    if template_data.subcategory_id:
        sub = await sub_service.get_subcategory(conn, template_data.subcategory_id)
        if sub: subcategory_name = sub["name"]

    render_data = template_data.model_dump()
    render_data.update({"industry_name": industry_name, "category_name": category_name, "subcategory_name": subcategory_name})

    docx_bytes = await generate_docx_template(render_data)
    unique_id = uuid.uuid4()
    # Ensure template_name is not None
    name_for_path = template_data.template_name or "template"
    s3_path = f"templates/{unique_id}_{name_for_path.replace(' ', '_')}.docx"
    s3_url = await upload_bytes_to_s3(docx_bytes, s3_path, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    now = datetime.now()
    data = template_data.model_dump()
    data.pop("prompt", None) # Remove helper field
    data.update({
        "category": category_name,
        "subcategory": subcategory_name,
        "tenant_id": current_user.tenant_id,
        "user_id": current_user.user_id,
        "header_image": s3_url,
        "created_on": now,
        "updated_on": now
    })
    
    result = await service.create_template(conn, data)
    
    # Record Audit Log
    await AuditLogDBService.record_audit_log(
        conn, current_user.tenant_id, current_user.user_id,
        "template_create", "template", str(result["template_id"]),
        {"template_name": result["template_name"]}
    )
    
    return result


@router.post("/{template_id}/generate")
async def generate_document_from_template(
    template_id: uuid.UUID,
    request: TemplateGenerateRequest,
    conn: psycopg.AsyncConnection = Depends(get_raw_db),
    current_user: Any = Depends(get_current_user),
    service: TemplateDBService = Depends(get_template_service)
):
    """
    Generate a document from a selected template skeleton by injecting data.
    Supports JSON and CSV data sources.
    """
    template = await service.get_template(conn, template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    # 1. Prepare Data
    render_data = request.data or {}
    
    # 2. Handle CSV Injection
    if request.csv_data:
        csv_rows = parse_csv_data(request.csv_data)
        render_data["rows"] = csv_rows
        render_data["csv"] = csv_rows # Support both names

    # 3. Handle Sections
    if request.sections:
        # Merge section-specific data into main render_data
        render_data.update(request.sections)

    # 4. Render HTML Content
    html_content = template.get("html_content")
    if not html_content:
        raise HTTPException(status_code=400, detail="Template has no HTML content skeleton")

    rendered_html = await render_html_template(html_content, render_data)

    # 5. Output Format Handling
    if request.output_format == "docx":
        # Create a document for DOCX export
        from docx import Document as DocxDocument
        import io
        doc = DocxDocument()
        doc.add_heading(template.get("template_name"), 0)
        # Simple HTML to text conversion for now, or use a better library if needed
        # In a real app, we'd use something like htmldocx
        doc.add_paragraph(rendered_html) 
        
        target_stream = io.BytesIO()
        doc.save(target_stream)
        file_bytes = target_stream.getvalue()
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    else:
        # Default to PDF (Mocking PDF generation for now, as we don't have a PDF lib installed)
        # We can use something like WeasyPrint or pdfkit in a real setup
        file_bytes = rendered_html.encode('utf-8')
        mime_type = "application/pdf" # Mocking MIME
        ext = "pdf"

    # 6. Upload to S3 and Save as a Document
    unique_id = uuid.uuid4()
    filename = request.filename or f"Generated_{template.get('template_name')}_{unique_id}.{ext}"
    s3_path = f"tenant-{current_user.tenant_id}/generated/{unique_id}_{filename}"
    
    s3_url = await upload_bytes_to_s3(file_bytes, s3_path, mime_type)

    # 7. Create Document record
    from app.services.db.document_db_service import DocumentDBService
    doc_service = DocumentDBService()
    
    doc_record = await doc_service.create_document(
        conn, unique_id, current_user.tenant_id, current_user.user_id,
        filename, s3_url, len(file_bytes)
    )
    
    # 8. Save initial version with the rendered content
    await doc_service.save_document_version(
        conn, unique_id, 1, rendered_html, json.dumps(render_data), current_user.user_id
    )

    # Record Audit Log
    await AuditLogDBService.record_audit_log(
        conn, current_user.tenant_id, current_user.user_id,
        "document_generate", "template", str(template_id),
        {"filename": filename, "doc_id": str(unique_id)}
    )

    return {
        "document_id": unique_id,
        "filename": filename,
        "url": s3_url,
        "status": "generated"
    }


@router.get("/{template_id}", response_model=TemplateResponse)
async def get_template(
    template_id: uuid.UUID, 
    conn: psycopg.AsyncConnection = Depends(get_raw_db),
    service: TemplateDBService = Depends(get_template_service)
):
    template = await service.get_template(conn, template_id)
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    return template


@router.patch("/{template_id}", response_model=TemplateResponse)
async def update_template(
    template_id: uuid.UUID, 
    template_update: TemplateUpdate, 
    conn: psycopg.AsyncConnection = Depends(get_raw_db),
    current_user: Any = Depends(get_current_user),
    service: TemplateDBService = Depends(get_template_service)
):
    existing = await service.get_template(conn, template_id)
    if not existing:
        raise HTTPException(status_code=404, detail="Template not found")
    
    if not existing.get("is_public") and existing.get("tenant_id") != current_user.tenant_id:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    update_data = template_update.model_dump(exclude_unset=True)
    result = await service.update_template(conn, template_id, update_data)
    
    # Record Audit Log
    await AuditLogDBService.record_audit_log(
        conn, current_user.tenant_id, current_user.user_id,
        "template_update", "template", str(template_id)
    )
    
    return result


@router.delete("/{template_id}")
async def delete_template(
    template_id: uuid.UUID, 
    conn: psycopg.AsyncConnection = Depends(get_raw_db),
    current_user: Any = Depends(get_current_user),
    service: TemplateDBService = Depends(get_template_service)
):
    deleted = await service.delete_template(conn, template_id, current_user.tenant_id)
    if not deleted:
         # Either not found or not authorized to delete
         raise HTTPException(status_code=404, detail="Template not found or unauthorized")
         
    # Record Audit Log
    await AuditLogDBService.record_audit_log(
        conn, current_user.tenant_id, current_user.user_id,
        "template_delete", "template", str(template_id)
    )
         
    return {"status": "deleted"}
