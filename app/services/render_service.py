from jinja2 import Template as JinjaTemplate
from typing import Any, Dict, List
import io
import csv
from docx import Document

def parse_csv_data(csv_string: str) -> List[Dict[str, Any]]:
    """
    Parses a CSV string into a list of dictionaries.
    """
    f = io.StringIO(csv_string)
    reader = csv.DictReader(f)
    return list(reader)

async def render_html_template(html_content: str, data: Dict[str, Any]) -> str:
    """
    Renders a rich HTML template using Jinja2 and the provided structured data.
    """
    if not html_content:
        return ""
    
    template = JinjaTemplate(html_content)
    return template.render(**data)

async def generate_docx_template(template_data: Dict[str, Any]) -> bytes:
    """
    Generates a professional .docx file based on template metadata.
    """
    doc = Document()
    
    # Header Section
    title = template_data.get("template_name", "Document Template")
    doc.add_heading(title, 0)
    
    # Metadata Table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Field'
    hdr_cells[1].text = 'Value'
    
    metadata = {
        "Industry": template_data.get("industry_name", "N/A"),
        "Category": template_data.get("category_name", "N/A"),
        "Subcategory": template_data.get("subcategory_name", "N/A"),
        "Description": template_data.get("description", "N/A"),
    }
    
    for key, value in metadata.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value)

    doc.add_paragraph("\n")
    
    # Main Content
    doc.add_heading('Template Structure', level=1)
    doc.add_paragraph(template_data.get("html_content", "Generic document structure placeholder."))
    
    # Footer
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = f"SaaS Document Platform - {template_data.get('template_name')} Template"
    
    # Save to BytesIO
    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream.getvalue()
